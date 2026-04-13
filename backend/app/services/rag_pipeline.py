"""RAG (Retrieval-Augmented Generation) pipeline service.

Three public entry points:
    scrape_website(tenant_id, url)              — crawl4ai scrape → chunk → embed → store
    upload_document(tenant_id, file_content, filename) — extract → chunk → embed → store
    query(tenant_id, question, top_k=5)         — embed query → cosine search → list[str]

All DB operations are scoped by tenant_id (multi-tenancy requirement).
"""

import asyncio
import logging
import re
import uuid
from io import BytesIO
from typing import TYPE_CHECKING

from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.rag_document import RagDocument
from app.services.embeddings import generate_embedding_async as generate_embedding

if TYPE_CHECKING:
    from app.models.tenant import Tenant

logger = logging.getLogger(__name__)

_CHUNK_SIZE_TOKENS = 500       # approximate token limit per chunk
_CHUNK_OVERLAP_CHARS = 100     # character overlap between consecutive chunks
_AVG_CHARS_PER_TOKEN = 4       # rough estimate for splitting without a tokeniser


# ── Public API ────────────────────────────────────────────────────────────────

async def scrape_website(
    db: AsyncSession,
    tenant_id: uuid.UUID,
    url: str,
    tenant: "Tenant | None" = None,
) -> list[RagDocument]:
    """Crawl *url*, chunk the extracted text, embed each chunk, and persist to rag_documents.

    Uses crawl4ai for async web crawling when available; falls back to httpx +
    BeautifulSoup for simpler pages.

    Args:
        db: Async database session.
        tenant_id: Owning tenant — all rows scoped to this ID.
        url: Starting URL to crawl (up to 20 pages followed via internal links).
        tenant: Optional Tenant for resolving the embedding API key.

    Returns:
        List of persisted RagDocument instances (one per chunk).
    """
    pages = await _crawl(url)
    documents: list[RagDocument] = []

    for page_url, page_text in pages:
        if not page_text.strip():
            continue
        chunks = _chunk_text(page_text)
        for chunk in chunks:
            doc = await _store_chunk(
                db=db,
                tenant_id=tenant_id,
                content_text=chunk,
                source_type="website_scrape",
                source_url=page_url,
                tenant=tenant,
            )
            documents.append(doc)

    logger.info(
        "scrape_website: tenant=%s url=%s pages=%d chunks=%d",
        tenant_id,
        url,
        len(pages),
        len(documents),
    )
    return documents


async def upload_document(
    db: AsyncSession,
    tenant_id: uuid.UUID,
    file_content: bytes,
    filename: str,
    tenant: "Tenant | None" = None,
) -> list[RagDocument]:
    """Extract text from an uploaded file, chunk, embed, and store in rag_documents.

    Supported formats: PDF (.pdf), DOCX (.docx), plain text (.txt).

    Args:
        db: Async database session.
        tenant_id: Owning tenant.
        file_content: Raw bytes of the uploaded file.
        filename: Original filename (used to determine format and for display).
        tenant: Optional Tenant for resolving the embedding API key.

    Returns:
        List of persisted RagDocument instances (one per chunk).
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    text = _extract_text(file_content, ext, filename)

    if not text.strip():
        logger.warning("upload_document: no text extracted from %r", filename)
        return []

    chunks = _chunk_text(text)
    documents: list[RagDocument] = []

    for chunk in chunks:
        doc = await _store_chunk(
            db=db,
            tenant_id=tenant_id,
            content_text=chunk,
            source_type="manual_upload",
            filename=filename,
            tenant=tenant,
        )
        documents.append(doc)

    logger.info(
        "upload_document: tenant=%s file=%r chunks=%d",
        tenant_id,
        filename,
        len(documents),
    )
    return documents


async def query(
    db: AsyncSession,
    tenant_id: uuid.UUID,
    question: str,
    top_k: int = 5,
    tenant: "Tenant | None" = None,
) -> list[str]:
    """Retrieve the top-k most relevant text chunks for *question*.

    Embeds the question and performs cosine similarity search against
    rag_documents using pgvector's <=> operator.

    Args:
        db: Async database session.
        tenant_id: Owning tenant — only that tenant's documents are searched.
        question: Natural-language query from the chat widget.
        top_k: Maximum number of chunks to return (default 5).
        tenant: Optional Tenant for resolving the embedding API key.

    Returns:
        List of content_text strings, ordered by relevance (most relevant first).
    """
    question_embedding = await generate_embedding(question, tenant)
    embedding_str = "[" + ",".join(str(v) for v in question_embedding) + "]"

    # pgvector cosine distance operator: <=>  (lower = more similar)
    sql = text(
        """
        SELECT content_text
        FROM rag_documents
        WHERE tenant_id = :tenant_id
        ORDER BY embedding <=> CAST(:embedding AS vector)
        LIMIT :top_k
        """
    )
    result = await db.execute(
        sql,
        {"tenant_id": str(tenant_id), "embedding": embedding_str, "top_k": top_k},
    )
    rows = result.fetchall()
    return [row[0] for row in rows]


# ── Internal helpers ──────────────────────────────────────────────────────────

async def _crawl(url: str) -> list[tuple[str, str]]:
    """Crawl *url* and return a list of (page_url, page_text) pairs.

    Attempts to use crawl4ai; falls back to httpx + BeautifulSoup.
    A 30-second timeout guards against crawl4ai hanging when Playwright
    cannot launch a browser (e.g. WSL2 without display).
    """
    try:
        return await asyncio.wait_for(_crawl_with_crawl4ai(url), timeout=30.0)
    except asyncio.TimeoutError:
        logger.warning("crawl4ai timed out after 30 s — falling back to httpx+BeautifulSoup")
    except ImportError:
        logger.debug("crawl4ai not installed — falling back to httpx+BeautifulSoup")
    except Exception as exc:
        logger.warning("crawl4ai failed (%s) — falling back to httpx+BeautifulSoup", exc)

    return await _crawl_with_httpx(url)


async def _crawl_with_crawl4ai(url: str) -> list[tuple[str, str]]:
    """Crawl using crawl4ai (async, JS-rendered)."""
    from crawl4ai import AsyncWebCrawler  # type: ignore[import]

    pages: list[tuple[str, str]] = []
    async with AsyncWebCrawler(verbose=False) as crawler:
        result = await crawler.arun(url=url)
        if result.success and result.markdown:
            pages.append((url, result.markdown))
    return pages


async def _crawl_with_httpx(url: str) -> list[tuple[str, str]]:
    """Fallback crawler: httpx + BeautifulSoup, follows up to 20 internal links."""
    import httpx
    from bs4 import BeautifulSoup
    from urllib.parse import urljoin, urlparse

    base_netloc = urlparse(url).netloc
    visited: set[str] = set()
    queue = [url]
    pages: list[tuple[str, str]] = []
    max_pages = 20

    async with httpx.AsyncClient(follow_redirects=True, timeout=15) as client:
        while queue and len(pages) < max_pages:
            current_url = queue.pop(0)
            if current_url in visited:
                continue
            visited.add(current_url)

            try:
                resp = await client.get(current_url)
                if resp.status_code != 200:
                    continue
                ct = resp.headers.get("content-type", "")
                if "html" not in ct:
                    continue

                soup = BeautifulSoup(resp.text, "html.parser")

                # Remove navigation, scripts, styles, footers.
                for tag in soup(["script", "style", "nav", "footer", "header"]):
                    tag.decompose()

                page_text = soup.get_text(separator="\n", strip=True)
                if page_text.strip():
                    pages.append((current_url, page_text))

                # Collect internal links.
                for anchor in soup.find_all("a", href=True):
                    href = urljoin(current_url, anchor["href"])
                    if urlparse(href).netloc == base_netloc and href not in visited:
                        queue.append(href)

            except Exception as exc:
                logger.debug("_crawl_with_httpx: error fetching %s: %s", current_url, exc)

    return pages


def _extract_text(file_content: bytes, ext: str, filename: str) -> str:
    """Extract plain text from PDF, DOCX, or TXT file bytes."""
    if ext == "pdf":
        return _extract_pdf(file_content)
    if ext == "docx":
        return _extract_docx(file_content)
    if ext == "txt":
        return file_content.decode("utf-8", errors="replace")

    logger.warning("_extract_text: unsupported file type %r for %r", ext, filename)
    return ""


def _extract_pdf(file_content: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    try:
        import pdfplumber  # type: ignore[import]

        with pdfplumber.open(BytesIO(file_content)) as pdf:
            parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text)
        return "\n".join(parts)
    except Exception as exc:
        logger.error("_extract_pdf: extraction failed: %s", exc)
        return ""


def _extract_docx(file_content: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        import docx  # type: ignore[import]

        doc = docx.Document(BytesIO(file_content))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as exc:
        logger.error("_extract_docx: extraction failed: %s", exc)
        return ""


def _chunk_text(text: str) -> list[str]:
    """Split *text* into overlapping chunks of approximately _CHUNK_SIZE_TOKENS tokens.

    Uses a simple character-count heuristic (_AVG_CHARS_PER_TOKEN) to avoid
    importing a full tokeniser.  Splits on paragraph boundaries when possible
    to preserve semantic coherence.
    """
    max_chars = _CHUNK_SIZE_TOKENS * _AVG_CHARS_PER_TOKEN
    # Normalise whitespace.
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    # Split into paragraphs first.
    paragraphs = [p.strip() for p in re.split(r"\n\n+", text) if p.strip()]

    chunks: list[str] = []
    current_chunk: list[str] = []
    current_len = 0

    for para in paragraphs:
        para_len = len(para)

        if current_len + para_len > max_chars and current_chunk:
            chunks.append("\n\n".join(current_chunk))
            # Overlap: carry last paragraph into next chunk.
            current_chunk = current_chunk[-1:] if current_chunk else []
            current_len = len(current_chunk[0]) if current_chunk else 0

        # Paragraph itself exceeds max — hard-split.
        if para_len > max_chars:
            for i in range(0, para_len, max_chars - _CHUNK_OVERLAP_CHARS):
                sub = para[i : i + max_chars]
                if sub.strip():
                    chunks.append(sub)
            current_chunk = []
            current_len = 0
        else:
            current_chunk.append(para)
            current_len += para_len

    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    return [c for c in chunks if c.strip()]


async def _store_chunk(
    db: AsyncSession,
    tenant_id: uuid.UUID,
    content_text: str,
    source_type: str,
    tenant: "Tenant | None" = None,
    source_url: str | None = None,
    filename: str | None = None,
) -> RagDocument:
    """Embed *content_text* and persist a RagDocument row."""
    embedding = await generate_embedding(content_text, tenant)

    doc = RagDocument(
        id=uuid.uuid4(),
        tenant_id=tenant_id,
        source_type=source_type,
        source_url=source_url,
        filename=filename,
        content_text=content_text,
        embedding=embedding,
    )
    db.add(doc)
    await db.flush()
    await db.commit()

    return doc
